import streamlit as st
import os
import PyPDF2
from openai import OpenAI
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor

# --- SIDEBAR WITH LOGO AND COMPANY INFO ---
with st.sidebar:
    st.image("logo.png", width=120)
    st.markdown("<h3 style='color:#06038D;'>ECR SYSTEM</h3>", unsafe_allow_html=True)
    st.markdown("<span style='color:#FF671F;'>Powered by SSO Consultants</span>", unsafe_allow_html=True)

# --- HEADER WITH LOGO ---
col1, col2 = st.columns([8, 1])
with col1:
    st.markdown("<h1 style='color:#06038D;'>ECR SYSTEM</h1>", unsafe_allow_html=True)
with col2:
    st.image("logo.png", width=80)

user_name = st.text_input("Enter your name")
if user_name:
    st.success(f"Welcome, {user_name}!")

st.markdown(
    "<h4 style='color:#FF671F;'>Upload your documents (application, proposal, questionnaire, etc.):</h4>",
    unsafe_allow_html=True
)
uploaded_files = st.file_uploader(
    "Upload PDF or TXT files",
    type=["pdf", "txt"],
    accept_multiple_files=True
)

run_review = st.button("Run Ethics Review")

# --- HELPER FUNCTIONS ---

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() or ""
    return text

def classify_user_doc(text):
    text_lower = text.lower()
    if "informed consent" in text_lower:
        return "Informed Consent"
    elif "questionnaire" in text_lower or "survey" in text_lower:
        return "Questionnaire"
    elif "research proposal" in text_lower or "introduction" in text_lower:
        return "Research Proposal"
    elif "application" in text_lower:
        return "Application Form"
    elif "rac" in text_lower and "confirmation" in text_lower:
        return "RAC Confirmation Letter"
    else:
        return "Unknown"

def parse_markdown_table(md_table):
    lines = [line.strip() for line in md_table.strip().split('\n') if line.strip()]
    if len(lines) < 2 or "|" not in lines[0]:
        return [], []
    headers = [h.strip() for h in lines[0].split("|") if h.strip()]
    rows = []
    for line in lines[2:]:
        if "|" in line:
            row = [c.strip() for c in line.split("|") if c.strip()]
            if row:
                rows.append(row)
    return headers, rows

def extract_section(text, section_number):
    import re
    pattern = rf"\n{section_number}\s*(.*?)(?=\n\d+\.\s|$)"
    match = re.search(pattern, text, re.DOTALL)
    return match.group(1).strip() if match else ""

def extract_first_table(text):
    import re
    table_pattern = r"(\|.+\|\n\|[-\s|:]+\|\n(?:\|.*\|\n?)+)"
    match = re.search(table_pattern, text)
    return match.group(1).strip() if match else ""

def get_summary_section(ai_review):
    import re
    summary_pattern = r"(Summary|Highlights|Key Findings)[:\n]+(.+?)(?=\n\d+\.\s|$)"
    match = re.search(summary_pattern, ai_review, re.IGNORECASE | re.DOTALL)
    if match:
        return match.group(2).strip()
    lines = ai_review.strip().split("\n")
    return "\n".join(lines[:4]) if lines else "No summary provided."

def create_docx_report(user_name, summary, ai_review, logo_path="logo.png"):
    doc = Document()

    # --- COVER/TITLE ROW WITH LOGO ---
    table = doc.add_table(rows=1, cols=2)
    cell_title = table.cell(0, 0)
    cell_logo = table.cell(0, 1)
    cell_title.text = "ECR Report"
    cell_title.paragraphs[0].runs[0].font.size = Pt(22)
    cell_title.paragraphs[0].runs[0].font.bold = True
    try:
        cell_logo.paragraphs[0].add_run().add_picture(logo_path, width=Inches(1.1))
    except Exception:
        pass
    doc.add_paragraph(f"Prepared for: {user_name}", style='Intense Quote')
    doc.add_paragraph("")

    # --- SUMMARY SECTION ---
    doc.add_heading("Summary", level=1)
    summary_text = get_summary_section(ai_review)
    p = doc.add_paragraph(summary_text)
    p.runs[0].font.color.rgb = RGBColor(255, 103, 31)  # saffron/orange
    doc.add_paragraph("")

    # --- USER DOCUMENT CLASSIFICATION TABLE ---
    doc.add_heading("User Document Classification Summary", level=1)
    if summary:
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Expected Type'
        hdr_cells[1].text = 'Detected In'
        hdr_cells[2].text = 'Status'
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(6, 3, 141)  # blue
        for row in summary:
            row_cells = table.add_row().cells
            row_cells[0].text = row['Expected Type']
            row_cells[1].text = row['Detected In']
            row_cells[2].text = row['Status']
    else:
        doc.add_paragraph("No document summary available.")
    doc.add_paragraph("")

    # --- EXTRACT AND RENDER AI TABLES ---
    # 1. Required Documents Table
    doc.add_heading("Required Documents Table", level=1)
    required_table_md = extract_first_table(extract_section(ai_review, "1."))
    if required_table_md:
        headers, rows = parse_markdown_table(required_table_md)
        if headers and rows:
            table = doc.add_table(rows=1, cols=len(headers))
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(6, 3, 141)
            for row in rows:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = val
        else:
            doc.add_paragraph("No required documents table found.")
    else:
        doc.add_paragraph("No required documents table found.")
    doc.add_paragraph("")

    # 2. Concerns & Explanation Table (from section 3)
    doc.add_heading("Questionnaire English & Construction Concerns", level=1)
    concerns_table_md = extract_first_table(extract_section(ai_review, "3."))
    if concerns_table_md:
        headers, rows = parse_markdown_table(concerns_table_md)
        if headers and rows:
            table = doc.add_table(rows=1, cols=len(headers))
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(6, 3, 141)
            for row in rows:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = val
        else:
            doc.add_paragraph("No concerns table found.")
    else:
        doc.add_paragraph("No concerns table found.")
    doc.add_paragraph("")

    # --- REMAINING SECTIONS (2, 4, 5, 6) ---
    for section_num, section_title in [
        ("2.", "Ethics Compliance"),
        ("4.", "Alignment Check"),
        ("5.", "Other Aspects"),
        ("6.", "Overall Recommendation"),
    ]:
        doc.add_heading(section_title, level=1)
        section_text = extract_section(ai_review, section_num)
        doc.add_paragraph(section_text if section_text else "No information provided.")
        doc.add_paragraph("")

    # --- FOOTER ---
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "©copyright SSO Consultants"
    footer_para.alignment = 1  # Center

    # --- STYLE ---
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)
    return doc

# --- MAIN LOGIC ---

if run_review and uploaded_files and user_name:
    with st.spinner("Processing your documents and submitting to GPT..."):
        # --- LOAD AND PROCESS REFERENCE DOCS ---
        REFERENCE_DOCS_PATH = "REFRENCE DOCS"
        reference_docs = []
        if os.path.exists(REFERENCE_DOCS_PATH):
            for f in os.listdir(REFERENCE_DOCS_PATH):
                if f.lower().endswith(('.pdf', '.txt')):
                    file_path = os.path.join(REFERENCE_DOCS_PATH, f)
                    if f.lower().endswith('.pdf'):
                        with open(file_path, "rb") as file:
                            text = extract_text_from_pdf(file)
                    else:
                        with open(file_path, "r", encoding="utf-8") as file:
                            text = file.read()
                    reference_docs.append({
                        "filename": f,
                        "text": text
                    })

        # --- PROCESS USER UPLOADED DOCS ---
        required_types = [
            "Application Form",
            "Research Proposal",
            "Questionnaire",
        ]
        optional_types = [
            "Informed Consent",
            "RAC Confirmation Letter"
        ]
        all_types = required_types + optional_types

        user_docs = []
        for file in uploaded_files:
            if file.name.lower().endswith(".pdf"):
                text = extract_text_from_pdf(file)
            elif file.name.lower().endswith(".txt"):
                text = file.read().decode("utf-8")
            else:
                text = ""
            doc_type = classify_user_doc(text)
            user_docs.append({
                "filename": file.name,
                "type_detected": doc_type,
                "text": text
            })

        # --- DOCUMENT SUMMARY TABLE ---
        summary = []
        for t in all_types:
            found = False
            for doc in user_docs:
                if doc['type_detected'] == t:
                    found = True
                    summary.append({
                        "Expected Type": t,
                        "Detected In": doc['filename'],
                        "Status": "OK"
                    })
            if not found:
                summary.append({
                    "Expected Type": t,
                    "Detected In": "",
                    "Status": "MISSING" if t in required_types else "Optional - Not Uploaded"
                })

        df = pd.DataFrame(summary)
        st.subheader("User Document Classification Summary")
        st.dataframe(df)

        # --- PREPARE PROMPT FOR OPENAI ---
        user_documents_text = ""
        for doc in user_docs:
            user_documents_text += f"{doc['type_detected']} ({doc['filename']}):\n{doc['text'][:1500]}\n\n"

        reference_documents_text = ""
        for doc in reference_docs:
            reference_documents_text += f"{doc['filename']}:\n{doc['text'][:1500]}\n\n"

        ethics_prompt = f"""
        You are an expert in India-related ethics committee working. You will provide answers based only on the reference documents provided below, except for English and grammar, where you may use your own expertise or other references.

        If any required user document is missing, or if a document appears to be mislabeled (e.g., the content of a file uploaded as "Questionnaire" looks like a "Research Proposal"), please point this out clearly at the beginning of your response, and suggest to the user which document(s) need to be uploaded or corrected.

        Your task is to review the following uploaded documents and provide a detailed analysis as per these points:

        1. Are all the required documents provided? Present this in a tabular format.
        2. Does this proposal meet ethics requirements as per the reference documents? Give section-wise compliance or non-compliance, with explanations where it is non-compliant.
        3. Compare the English and construction of the questionnaire and highlight any concerns (present in a table).
        4. Does the questionnaire and informed consent align with the research proposal?
        5. Any other aspect that you might want to highlight.
        6. Overall recommendation and questions to ask (and why).

        Please start your answer with a 2-3 line summary of your findings.

        User Name:
        {user_name}

        User Documents:
        {user_documents_text}

        Reference Documents:
        {reference_documents_text}
        """

        # --- CALL OPENAI ---
        client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": ethics_prompt}],
            max_tokens=1500,
            temperature=0.2,
        )
        ai_review = response.choices[0].message.content

        st.subheader("📄 ECR Report")
        st.write(f"Hello {user_name}, here is your review:")
        st.write(ai_review)

        # --- CREATE STRUCTURED DOCX REPORT ---
        doc = create_docx_report(user_name, summary, ai_review)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        file_name = f"{user_name.replace(' ', '_')}_ECR_Report.docx"
        st.download_button(
            label="Download ECR Report as Word (.docx)",
            data=bio,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=file_name
        )

elif run_review and not user_name:
    st.warning("Please enter your name above to proceed.")
elif run_review and not uploaded_files:
    st.warning("Please upload at least one document to proceed.")

# --- FOOTER ---
st.markdown(
    '<div style="text-align:center; color:#FF671F; margin-top:30px;">©copyright SSO Consultants</div>',
    unsafe_allow_html=True
)
